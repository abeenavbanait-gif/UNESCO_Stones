import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx_helpers import set_cell_background, set_cell_margins, add_callout, style_heading, add_body_p
import os

doc = docx.Document()

# Page setup: Standard 1 inch margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Title Block
p_title = doc.add_paragraph()
p_title.paragraph_format.space_before = Pt(0)
p_title.paragraph_format.space_after = Pt(4)
r_title = p_title.add_run("HERITAGE STONES PROJECT: METHODOLOGICAL SCOPE & DATA LIMITATIONS REPORT")
r_title.font.name = 'Arial'
r_title.font.size = Pt(20)
r_title.bold = True
r_title.font.color.rgb = RGBColor(197, 48, 48) # Dark Red for Critical Risk

p_sub = doc.add_paragraph()
p_sub.paragraph_format.space_after = Pt(14)
r_sub = p_sub.add_run("An Analytical Audit of Extraction Gaps, Automated NLP/LLM Failure Modes, and OUV Document Constraints (Backup 34 Audit)")
r_sub.font.name = 'Arial'
r_sub.font.size = Pt(11.5)
r_sub.font.color.rgb = RGBColor(74, 85, 104)

# Metadata Table
meta_table = doc.add_table(rows=2, cols=4)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["Audit Scope", "Unstudied Site Volume", "Identified OUV Issues", "Estimated Hidden Mass"]
vals = ["Backup 34 (July 2026)", "709 Sites (78.6%)", "22 Issue / 24 Absent", "~532 Uncaptured Stone Sites"]

for i in range(4):
    c_hdr = meta_table.cell(0, i)
    set_cell_background(c_hdr, "742A2A")
    set_cell_margins(c_hdr, top=80, bottom=80, left=100, right=100)
    p = c_hdr.paragraphs[0]
    r = p.add_run(headers[i])
    r.font.name = 'Arial'
    r.font.size = Pt(9)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)
    
    c_val = meta_table.cell(1, i)
    set_cell_background(c_val, "FFF5F5")
    set_cell_margins(c_val, top=80, bottom=80, left=100, right=100)
    p2 = c_val.paragraphs[0]
    r2 = p2.add_run(vals[i])
    r2.font.name = 'Arial'
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(45, 55, 72)

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# SECTION 1: EXECUTIVE METHODOLOGICAL AUDIT
style_heading(doc.add_paragraph(), "1. Executive Audit & Scope Definition", level=1)

add_body_p(doc, "Every empirical dataset carries methodological boundaries. While the analytical report summarizes findings from 193 verified UNESCO sites, this companion audit examines the remaining 709 sites (78.6% of the 902-record universe) currently lacking structured stone entries. Following the Backup 34 audit, 6 spurious 'Flint' data entries were successfully scrubbed, and 6 additional sites with inaccessible OUV documents (including Dubrovnik and Selimiye Mosque) were formally flagged as OUV Absent.")

add_callout(doc, 
    "Out of 709 unstudied sites, an estimated 532 sites (~75%) actually utilize geological building materials extensively. Automated NLP/LLM scrapers fail to capture these sites primarily because official UNESCO Outstanding Universal Value (OUV) texts focus on historical significance rather than material composition. Over 200 sites feature implicit stone construction (e.g., Gothic cathedrals, ancient amphitheatres) where stone is architecturally present but never explicitly named in the official text.",
    title="METHODOLOGICAL AUDIT SUMMARY (BACKUP 34 AUDIT)", border_color="C53030", bg_color="FFF5F5")

# SECTION 2: THE 78% DATA GAP
style_heading(doc.add_paragraph(), "2. Anatomy of the 78% Data Gap", level=1)

add_body_p(doc, "The 709 unstudied sites are categorized using an internal operational classification system embedded within the dataset:")

# Status breakdown table
t_status = doc.add_table(rows=6, cols=3)
t_status.alignment = WD_TABLE_ALIGNMENT.CENTER
t_headers = ["Operational Category", "Site Count", "Methodological Explanation"]
t_data = [
    ["Skipped Research Queue", "649 sites (71.9%)", "Reviewed sites queued for subsequent manual archival extraction."],
    ["OUV Text Inaccessible (Absent)", "24 sites (2.7%)", "Official UNESCO OUV documentation is missing, corrupted, or unindexed (e.g. Dubrovnik, Selimiye Mosque, Ephesus)."],
    ["OUV Text Issues", "22 sites (2.4%)", "OUV text accessed but fails to mention stone despite obvious physical presence (e.g. Jerusalem, Göbekli Tepe)."],
    ["Architecture Context Only", "14 sites (1.6%)", "Basic architectural style recorded, but lithological entries pending."],
    ["Researched & Verified Baseline", "193 sites (21.4%)", "Fully completed baseline entries with citeable internal/external references."]
]

for col_idx, h in enumerate(t_headers):
    c = t_status.cell(0, col_idx)
    set_cell_background(c, "742A2A")
    set_cell_margins(c, top=80, bottom=80, left=100, right=100)
    p = c.paragraphs[0]
    r = p.add_run(h)
    r.font.name = 'Arial'
    r.font.size = Pt(9.5)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)

for row_idx, r_data in enumerate(t_data):
    bg = "FFFFFF" if row_idx % 2 == 0 else "FFF5F5"
    for col_idx, val in enumerate(r_data):
        c = t_status.cell(row_idx + 1, col_idx)
        set_cell_background(c, bg)
        set_cell_margins(c, top=70, bottom=70, left=100, right=100)
        p = c.paragraphs[0]
        r = p.add_run(val)
        r.font.name = 'Arial'
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(45, 55, 72)
        if col_idx == 1:
            r.bold = True

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# Add Image 8: Data Gap Donut
if os.path.exists('charts/data_gap_donut.png'):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after = Pt(4)
    p_img.add_run().add_picture('charts/data_gap_donut.png', width=Inches(5.2))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(14)
    r_cap = p_cap.add_run("Figure 1: Breakdown of researched vs. unstudied status across all 902 records (Backup 34 Audit).")
    r_cap.font.name = 'Arial'
    r_cap.font.size = Pt(8.5)
    r_cap.italic = True
    r_cap.font.color.rgb = RGBColor(113, 128, 150)

# SECTION 3: GEOGRAPHIC & NATIONAL GAPS
style_heading(doc.add_paragraph(), "3. Geographic Distribution of Data Gaps", level=1)

add_body_p(doc, "The research queue is not evenly distributed across nations. Unstudied sites cluster disproportionately in the world's most historic heritage nations: Italy (34 unstudied sites), Germany (33 sites), China (32 sites), France (31 sites), and Spain (30 sites).")

# Add Image 9: Country Gaps
if os.path.exists('charts/country_unstudied_gaps.png'):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after = Pt(4)
    p_img.add_run().add_picture('charts/country_unstudied_gaps.png', width=Inches(6.0))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(14)
    r_cap = p_cap.add_run("Figure 2: Comparison of unstudied site gaps against total inscribed sites for leading heritage nations.")
    r_cap.font.name = 'Arial'
    r_cap.font.size = Pt(8.5)
    r_cap.italic = True
    r_cap.font.color.rgb = RGBColor(113, 128, 150)

add_body_p(doc, "Because over 70% of inscribed monuments in Western Europe, East Asia, and Mesoamerica remain in the research queue, current dataset-wide statistics (such as the dominance of limestone or local provenance) reflect early research prioritization (in India, Jordan, UK) rather than the global equilibrium. Re-evaluating these metrics upon full queue completion will likely shift continental ratios.")

# SECTION 4: WHY NLP AND LLMS FAIL
style_heading(doc.add_paragraph(), "4. Methodological Analysis: Why NLP & LLMs Fail to Capture Heritage Stones", level=1)

add_body_p(doc, "A central question facing digital heritage initiatives is whether automated Natural Language Processing (NLP) or Large Language Model (LLM) pipelines can automatically extract stone types from official texts. Our empirical audit reveals six fundamental failure modes that render automated text extraction insufficient without human expert verification:")

failures = [
    ("1. Implicit vs. Explicit Material References",
     "UNESCO inscription texts focus on historical, cultural, and aesthetic significance rather than architectural materials. Official texts routinely describe 'Gothic cathedrals with flying buttresses' (e.g., Chartres, Amiens, Cologne) without ever explicitly writing the word 'limestone'. An automated NLP parser searching for explicit rock keywords registers a false negative, ignoring massive limestone structures."),
    
    ("2. Local & Vernacular Terminology Disconnect",
     "Regional quarrying traditions utilize hyper-local stone names that standard language models fail to map to geological classes. Terms like 'Sillar' (Peru ignimbrite), 'Aachener Blaustein' (German Devonian limestone), 'Pietra Serena' (Florentine greywacke), 'Daga' (Zimbabwe composite), and 'Kabook' (Sri Lankan laterite) are either misclassified or ignored by standard English NLP tokenizers."),
    
    ("3. Cross-Language Documentation Fragmentation",
     "Many official inscription dossiers exist primarily in French, Spanish, German, Arabic, or Chinese. English summaries provided by UNESCO routinely strip out specific material technical terms (e.g., translating precise French 'calcaire lumachelle' to generic 'local stone'). Automated extraction on English text alone suffers from severe translation loss."),
    
    ("4. Reliance on External Secondary Literature",
     "The most precise lithological data rarely exists within the brief UNESCO Statement of Outstanding Universal Value (OUV). For example, the precise geological identification of Arequipa's Sillar as 'Dacitic to Rhyolitic Vitric Tuff' or the Taj Mahal's marble as '>98% Calcite Marble' derives from specialized peer-reviewed petrographic papers, not UNESCO documents. Scrapers confined to OUV URLs miss the core scientific evidence."),
    
    ("5. Euphemistic & Architectural Vocabulary Obscuration",
     "Texts frequently utilize stylistic vocabulary ('ashlar masonry', 'opus quadratum', 'monolithic cella', 'rusticated façade') that implies worked stone without naming the specific rock type. NLP algorithms cannot reliably infer petrology from architectural style without domain-specific ontology rules."),
    
    ("6. The Closed-World Extraction Fallacy",
     "Automated models operate under a closed-world assumption: if 'limestone' is unmentioned, it is assumed absent. In cultural heritage documentation, silence regarding material composition reflects the committee's drafting priorities, not the physical absence of stone.")
]

for title, desc in failures:
    style_heading(doc.add_paragraph(), title, level=2)
    add_body_p(doc, desc)

# SECTION 5: THE ICEBERG MODEL & UNFILLED SITES
style_heading(doc.add_paragraph(), "5. The Iceberg Model & Hidden Stone Heritage", level=1)

add_body_p(doc, "The current dataset represents an 'Iceberg Model': the 193 verified stone sites form the visible tip above water, while an estimated 532 unstudied stone monuments lie beneath the surface, hidden by documentation gaps.")

# Add Image 10: Iceberg Model
if os.path.exists('charts/iceberg_uncaptured_model.png'):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after = Pt(4)
    p_img.add_run().add_picture('charts/iceberg_uncaptured_model.png', width=Inches(5.5))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(14)
    r_cap = p_cap.add_run("Figure 3: The Iceberg Model comparing documented verified sites against the estimated uncaptured stone universe.")
    r_cap.font.name = 'Arial'
    r_cap.font.size = Pt(8.5)
    r_cap.italic = True
    r_cap.font.color.rgb = RGBColor(113, 128, 150)

style_heading(doc.add_paragraph(), "High-Priority OUV Issue & Inaccessible Sites", level=2)
add_body_p(doc, "Our audit highlights 22 sites flagged with 'OUV Text Issues' and 24 sites flagged with 'Inaccessible OUV' (including newly audited sites such as Old City of Dubrovnik and Selimiye Mosque) that require immediate expert intervention. Crucially, these lists include some of the most famous stone monuments on Earth:")

# Priority sites table
t_priority = doc.add_table(rows=8, cols=3)
t_priority.alignment = WD_TABLE_ALIGNMENT.CENTER
t_p_headers = ["Monument / Site Name", "Location", "Expected Physical Lithology"]
t_p_data = [
    ["Old City of Dubrovnik", "Croatia", "Korčula and Brač Cretaceous limestone masonry walls and stone paving."],
    ["Selimiye Mosque Complex", "Türkiye", "Marmara marble columns, local limestone masonry, and granite piers."],
    ["Old City of Jerusalem & Walls", "Jerusalem", "Meleke Limestone ('Jerusalem Stone') - Cretaceous hard nari limestone."],
    ["Göbekli Tepe", "Türkiye", "Hand-carved megalithic T-pillars of local limestone (~11,600 BP)."],
    ["Historic Areas of Istanbul", "Türkiye", "Proconnesian Marble, Byzantine brick, and local Tertiary limestone."],
    ["Ephesus", "Türkiye", "Hellenistic and Roman marble temples, agoras, and paved avenues."],
    ["Historical Complex of Split", "Croatia", "Diocletian's Palace: Proconnesian marble and local Cretaceous limestone."]
]

for col_idx, h in enumerate(t_p_headers):
    c = t_priority.cell(0, col_idx)
    set_cell_background(c, "742A2A")
    set_cell_margins(c, top=80, bottom=80, left=100, right=100)
    p = c.paragraphs[0]
    r = p.add_run(h)
    r.font.name = 'Arial'
    r.font.size = Pt(9.5)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)

for row_idx, r_data in enumerate(t_p_data):
    bg = "FFFFFF" if row_idx % 2 == 0 else "FFF5F5"
    for col_idx, val in enumerate(r_data):
        c = t_priority.cell(row_idx + 1, col_idx)
        set_cell_background(c, bg)
        set_cell_margins(c, top=70, bottom=70, left=100, right=100)
        p = c.paragraphs[0]
        r = p.add_run(val)
        r.font.name = 'Arial'
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(45, 55, 72)
        if col_idx == 0:
            r.bold = True

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# SECTION 6: RECOMMENDATIONS
style_heading(doc.add_paragraph(), "6. Strategic Recommendations for Dataset Scaling", level=1)

add_body_p(doc, "To overcome these methodological boundaries and achieve a complete 902-site global inventory, the Heritage Stones project should implement a three-tiered roadmap:")

recs = [
    ("1. Deploy Hybrid AI-Expert Extraction Pipelines",
     "Replace naive keyword scrapers with a hybrid workflow. Use LLMs fine-tuned on architectural history to flag 'implicit stone' passages, followed by manual review by heritage geologists to confirm petrographic assignments."),

    ("2. Ingest Secondary ICOMOS & Academic Corpora",
     "Expand the ingestion pipeline beyond basic OUV summary text. Automatically parse ICOMOS Advisory Body Evaluations, State of Conservation reports, and peer-reviewed geology journals where technical petrographic details are recorded."),

    ("3. Implement Source Reliability & Confidence Metadata",
     "Introduce mandatory provenance tags for every dataset entry, distinguishing between 'Confirmed by OUV', 'Verified via Peer Literature', and 'Inferred from Architectural Style'. This maintains absolute scientific transparency as the dataset expands.")
]

for title, desc in recs:
    style_heading(doc.add_paragraph(), title, level=2)
    add_body_p(doc, desc)

# Save document
output_filename = "Heritage_Stones_Project_Limitations_and_Scope.docx"
doc.save(output_filename)
print(f"Report 2 successfully saved to {output_filename}")
