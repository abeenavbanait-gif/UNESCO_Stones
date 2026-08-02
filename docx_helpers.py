import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_callout(doc, text, title="KEY RESEARCH TAKEAWAY", border_color="1A365D", bg_color="F7FAFC"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24') # 3pt
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), border_color)
    tcBorders.append(left)
    
    for b in ['top', 'bottom', 'right']:
        node = OxmlElement(f'w:{b}')
        node.set(qn('w:val'), 'none')
        tcBorders.append(node)
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    
    r_title = p.add_run(f"{title}\n")
    r_title.bold = True
    r_title.font.name = 'Arial'
    r_title.font.size = Pt(10)
    r_title.font.color.rgb = RGBColor(26, 54, 93)
    
    r_text = p.add_run(text)
    r_text.font.name = 'Arial'
    r_text.font.size = Pt(9.5)
    r_text.font.color.rgb = RGBColor(45, 55, 72)
    
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(6)

def style_heading(p, text, level=1):
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.bold = True
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(26, 54, 93)
    elif level == 2:
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(43, 108, 176)
    elif level == 3:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(45, 55, 72)

def add_body_p(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(45, 55, 72)
    return p
