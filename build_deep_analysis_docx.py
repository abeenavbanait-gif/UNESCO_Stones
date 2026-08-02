import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx_helpers import set_cell_background, set_cell_margins, add_callout, style_heading, add_body_p
import os

doc = docx.Document()

# Page setup: Standard 1 inch margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Title Block
p_title = doc.add_paragraph()
p_title.paragraph_format.space_before = Pt(0)
p_title.paragraph_format.space_after = Pt(4)
r_title = p_title.add_run("GEOLOGICAL BIOGRAPHY OF WORLD HERITAGE")
r_title.font.name = 'Arial'
r_title.font.size = Pt(22)
r_title.bold = True
r_title.font.color.rgb = RGBColor(26, 54, 93)

p_sub = doc.add_paragraph()
p_sub.paragraph_format.space_after = Pt(14)
r_sub = p_sub.add_run("An Analytical Monograph on Lithology, Provenance, and Preservation Across UNESCO World Heritage Sites (Backup 34 Audit)")
r_sub.font.name = 'Arial'
r_sub.font.size = Pt(12)
r_sub.font.color.rgb = RGBColor(74, 85, 104)

# Metadata Table
meta_table = doc.add_table(rows=2, cols=4)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["Dataset Version", "Total Record Universe", "Geographic Scope", "Primary Source Basis"]
vals = ["Backup 34 (July 2026)", "902 Sites / 80 Attributes", "168 Sovereign Nations", "UNESCO OUV & Peer Literature"]

for i in range(4):
    c_hdr = meta_table.cell(0, i)
    set_cell_background(c_hdr, "1A365D")
    set_cell_margins(c_hdr, top=80, bottom=80, left=100, right=100)
    p = c_hdr.paragraphs[0]
    r = p.add_run(headers[i])
    r.font.name = 'Arial'
    r.font.size = Pt(9)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)
    
    c_val = meta_table.cell(1, i)
    set_cell_background(c_val, "F7FAFC")
    set_cell_margins(c_val, top=80, bottom=80, left=100, right=100)
    p2 = c_val.paragraphs[0]
    r2 = p2.add_run(vals[i])
    r2.font.name = 'Arial'
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(45, 55, 72)

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# SECTION 1: EXECUTIVE SUMMARY
style_heading(doc.add_paragraph(), "1. Executive Summary & Core Empirical Baseline", level=1)

add_body_p(doc, "The UNESCO World Heritage list represents the highest designation of human architectural and cultural achievement. Yet behind historical dates and cultural narratives lies a physical reality: every stone monument is a geological artifact, constructed from specific lithologies extracted from precise Earth formations. This updated analytical report presents findings from the latest Heritage Stones dataset (Backup 34), comprising 902 World Heritage Sites across 168 countries evaluating 80 distinct structural, lithological, and preservation attributes.")

add_callout(doc, 
    "Out of 902 total inscribed sites, exactly 193 entries (21.4%) contain verified primary stone identification data following the Backup 34 quality audit (which successfully cleaned up 6 erroneous 'Flint' data entries from non-flint sites). Among the 129 provenance-classified monuments, 88.4% relied exclusively on local stone extraction within a 5 to 25 km radius. Sedimentary formations (limestone and sandstone) account for 60.5% of all classified stone heritage, reflecting both the historical ease of manual masonry working and the widespread geographical availability of carbonate platforms near historic population centers.",
    title="EXECUTIVE RESEARCH BRIEF (BACKUP 34 AUDIT)", border_color="1A365D", bg_color="F7FAFC")

# Key Metrics Table
t_metrics = doc.add_table(rows=7, cols=3)
t_metrics.alignment = WD_TABLE_ALIGNMENT.CENTER
t_headers = ["Analytical Dimension", "Empirical Baseline", "Primary Heritage Significance"]
t_data = [
    ["Researched Stone Sites", "193 / 902 (21.4%)", "Verified research baseline; 709 sites form the unstudied expansion queue."],
    ["Dominant Rock Class", "Sedimentary (60.5%)", "Limestone (24 sites) and Sandstone (14 sites) anchor global masonry."],
    ["Resource Provenance", "Local (88.4%)", "114 of 129 provenance-classified sites extracted stone within immediate vicinity."],
    ["Strategic Sourcing (Imports)", "Imported (4.7%)", "Long-distance trade reserved for imperial prestige (Taj Mahal, CST Mumbai)."],
    ["Preservation Baseline", "78.4% Good/Excellent", "152 of 194 assessed sites exhibit strong structural integrity under UNESCO oversight."],
    ["Primary Degradation Vector", "Surface Erosion (31 sites)", "Rainwater leaching, wind abrasion, and humidity dominate structural threats."]
]

for col_idx, h in enumerate(t_headers):
    c = t_metrics.cell(0, col_idx)
    set_cell_background(c, "1A365D")
    set_cell_margins(c, top=80, bottom=80, left=100, right=100)
    p = c.paragraphs[0]
    r = p.add_run(h)
    r.font.name = 'Arial'
    r.font.size = Pt(9.5)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)

for row_idx, r_data in enumerate(t_data):
    bg = "FFFFFF" if row_idx % 2 == 0 else "F7FAFC"
    for col_idx, val in enumerate(r_data):
        c = t_metrics.cell(row_idx + 1, col_idx)
        set_cell_background(c, bg)
        set_cell_margins(c, top=70, bottom=70, left=100, right=100)
        p = c.paragraphs[0]
        r = p.add_run(val)
        r.font.name = 'Arial'
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(45, 55, 72)
        if col_idx == 1:
            r.bold = True

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# SECTION 2: DATASET ARCHITECTURE & COMPLETENESS
style_heading(doc.add_paragraph(), "2. Dataset Architecture & Empirical Completeness", level=1)

add_body_p(doc, "Building a rigorous geological inventory of global heritage requires granular field documentation. The 80 attributes in the dataset track three core operational dimensions: administrative metadata (Criteria, Country, Architecture Type), geological petrographic data (Rock Class, Lithology, Formation, Minerals, Colour, Texture), and architectural conservation metrics (Masonry Technique, Structural Use, Quarry Origin, Weathering, Condition, Restoration).")

add_body_p(doc, "Data completeness varies substantially across these layers. Administrative and geographical identifiers maintain a 100% completion rate, while detailed scientific indicators—such as specific mineralogy (2.4%) and stratigraphic formation (2.7%)—represent highly specialized entry points derived from peer-reviewed petrographic studies.")

# Add Image 5: Data Completeness
if os.path.exists('charts/data_completeness_bar.png'):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after = Pt(4)
    p_img.add_run().add_picture('charts/data_completeness_bar.png', width=Inches(6.0))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(14)
    r_cap = p_cap.add_run("Figure 1: Data completeness percentages by field across 902 UNESCO World Heritage Sites (Backup 34 Audit).")
    r_cap.font.name = 'Arial'
    r_cap.font.size = Pt(8.5)
    r_cap.italic = True
    r_cap.font.color.rgb = RGBColor(113, 128, 150)

# SECTION 3: GEOLOGICAL CLASSIFICATION
style_heading(doc.add_paragraph(), "3. Geological Classification: The Three Kingdoms of Stone", level=1)

add_body_p(doc, "Of the 152 sites with a confirmed rock classification, sedimentary geology dominates the historical landscape. Sedimentary rocks comprise 60.5% (92 sites), followed by igneous formations at 24.3% (37 sites), and metamorphic rocks at 15.1% (23 sites).")

# Add Image 1: Rock Class Donut
if os.path.exists('charts/rock_class_donut.png'):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after = Pt(4)
    p_img.add_run().add_picture('charts/rock_class_donut.png', width=Inches(5.2))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(14)
    r_cap = p_cap.add_run("Figure 2: Distribution of primary rock classes across classified heritage monuments (n=152).")
    r_cap.font.name = 'Arial'
    r_cap.font.size = Pt(8.5)
    r_cap.italic = True
    r_cap.font.color.rgb = RGBColor(113, 128, 150)

add_body_p(doc, "The pronounced dominance of sedimentary formations stems from three historical factors. First, limestones and sandstones possess superior workability under pre-industrial manual tools, allowing stone masons to achieve precise ashlar block geometry and fine relief carving. Second, shallow marine carbonate platforms cover extensive areas of Southern Europe, North Africa, and South Asia—precisely where early urban civilizations developed. Third, sedimentary rocks frequently form accessible surface bluffs and river valley outcrops, minimizing overburden excavation costs.")

style_heading(doc.add_paragraph(), "Regional Geological Patterns", level=2)
add_body_p(doc, "Geological utilization varies distinctly by continental landmass. Europe exhibits the highest concentration of classified sites (66 sites), led by sedimentary limestone cathedrals and classical Mediterranean marble structures. Asia demonstrates a strong dual reliance on sedimentary sandstones (Mughal and Harappan architecture) and Deccan basalt igneous formations (rock-cut cave temples). Africa and the Americas display balanced distributions, where volcanic tuffs and granites provided durable structural mediums for ancient civilizations.")

# Add Image 4: Regional Rock Class
if os.path.exists('charts/regional_rock_class.png'):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after = Pt(4)
    p_img.add_run().add_picture('charts/regional_rock_class.png', width=Inches(5.8))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(14)
    r_cap = p_cap.add_run("Figure 3: Rock class distribution broken down by geographic region.")
    r_cap.font.name = 'Arial'
    r_cap.font.size = Pt(8.5)
    r_cap.italic = True
    r_cap.font.color.rgb = RGBColor(113, 128, 150)

# SECTION 4: MAJOR STONES TAXONOMY
style_heading(doc.add_paragraph(), "4. Typological Dominance & Material Taxonomy", level=1)

add_body_p(doc, "Evaluating specific stone types reveals limestone as the single most frequently documented material in world heritage, appearing in 24 distinct major site entries, followed by granite (15 entries), sandstone (14 entries), and marble (10 entries). Following the Backup 34 audit, spurious 'Flint' assignments across 6 non-prehistoric sites were corrected, leaving Spiennes (Belgium) as the genuine Neolithic flint mining entry.")

# Add Image 2: Top Stones Bar Chart
if os.path.exists('charts/top_stones_bar.png'):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after = Pt(4)
    p_img.add_run().add_picture('charts/top_stones_bar.png', width=Inches(5.8))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(14)
    r_cap = p_cap.add_run("Figure 4: Frequency of primary stone types directly identified across UNESCO inscriptions (Backup 34 Audit).")
    r_cap.font.name = 'Arial'
    r_cap.font.size = Pt(8.5)
    r_cap.italic = True
    r_cap.font.color.rgb = RGBColor(113, 128, 150)

add_body_p(doc, "In addition to single-stone structures, several monuments showcase intentional multi-stone pairing. The most celebrated material pairing in the dataset is Red Sandstone combined with White Marble, characteristic of Mughal Imperial architecture in India (Taj Mahal, Humayun's Tomb, Fatehpur Sikri, Agra Fort). This dualism combined the structural durability and warm earthy presence of ferruginous sandstone with the spiritual purity and light reflection of pure calcitic marble.")

# SECTION 5: PROVENANCE & LOGISTICS
style_heading(doc.add_paragraph(), "5. Provenance Dynamics: Local Sourcing vs. Strategic Imports", level=1)

add_body_p(doc, "Overwhelmingly, historical builders utilized stone quarried within the immediate geographic vicinity. Among the 129 sites with recorded provenance data, 114 sites (88.4%) relied entirely on local extraction. Only 6 sites (4.7%) relied exclusively on long-distance stone imports, while 8 sites (6.2%) combined local structural stone with imported decorative elements.")

# Add Image 3: Provenance Pie
if os.path.exists('charts/provenance_pie.png'):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after = Pt(4)
    p_img.add_run().add_picture('charts/provenance_pie.png', width=Inches(5.2))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(14)
    r_cap = p_cap.add_run("Figure 5: Sourcing provenance distribution across documented heritage sites (n=129).")
    r_cap.font.name = 'Arial'
    r_cap.font.size = Pt(8.5)
    r_cap.italic = True
    r_cap.font.color.rgb = RGBColor(113, 128, 150)

add_body_p(doc, "The logistics of pre-industrial transport created severe physical constraints. Overland cart transport of heavy stone blocks over unpaved terrain imposed immense financial and labor costs. Consequently, long-distance stone movement occurred almost exclusively under two conditions: (1) maritime river and coastal shipping routes (such as Istrian stone transported across the Adriatic Sea to Venice), or (2) royal imperial decrees where cost was secondary to state symbolism (such as Makrana marble hauled 300+ km to Agra for the Taj Mahal, or Welsh bluestones transported 240 km to Stonehenge).")

# SECTION 6: CONSERVATION & WEATHERING
style_heading(doc.add_paragraph(), "6. Preservation Dynamics: Condition & Degradation Threat Vectors", level=1)

add_body_p(doc, "The dataset tracks environmental conservation status across 194 assessed sites. A strong majority—78.4% (152 sites)—are rated in Good or Excellent condition, testifying to active heritage protection and continuous site maintenance. However, 16.5% (32 sites) exhibit moderate structural degradation, and 5.2% (10 sites) suffer from poor or critical condition ratings requiring immediate conservation intervention.")

# Add Image 6: Condition Donut
if os.path.exists('charts/condition_assessment_pie.png'):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after = Pt(4)
    p_img.add_run().add_picture('charts/condition_assessment_pie.png', width=Inches(5.2))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(14)
    r_cap = p_cap.add_run("Figure 6: Environmental condition assessment distribution across evaluated monuments (n=194).")
    r_cap.font.name = 'Arial'
    r_cap.font.size = Pt(8.5)
    r_cap.italic = True
    r_cap.font.color.rgb = RGBColor(113, 128, 150)

style_heading(doc.add_paragraph(), "Deterioration Mechanisms", level=2)
add_body_p(doc, "Evaluating documented weathering threats across 119 sites highlights Surface Erosion (31 sites) and High Ambient Humidity (17 sites) as the two primary agents of material decay. Atmospheric pollution and sulfur dioxide emissions (11 sites) drive severe chemical weathering—specifically the sulfation of calcium carbonate in limestones and marbles into soluble gypsum crusts.")

# Add Image 7: Weathering Threats Bar
if os.path.exists('charts/weathering_threats_bar.png'):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after = Pt(4)
    p_img.add_run().add_picture('charts/weathering_threats_bar.png', width=Inches(5.8))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(14)
    r_cap = p_cap.add_run("Figure 7: Frequency of primary weathering vectors affecting heritage stone structures.")
    r_cap.font.name = 'Arial'
    r_cap.font.size = Pt(8.5)
    r_cap.italic = True
    r_cap.font.color.rgb = RGBColor(113, 128, 150)

# SECTION 7: CASE STUDIES
style_heading(doc.add_paragraph(), "7. Comprehensive Site Profiles: Benchmark Case Studies", level=1)

add_body_p(doc, "To demonstrate the depth of multi-attribute documentation in the dataset, the following eight case studies highlight monuments with exceptional data completeness (7 to 11 complete stone fields):")

case_studies = [
    ("Taj Mahal (Agra, India)", "Metamorphic Calcitic Marble (Makrana Marble) & Red Sandstone", 
     "Comprising 11 filled stone fields, the Taj Mahal represents the pinnacle of dual-material imperial masonry. The main mausoleum is clad in pure granoblastic calcitic marble quarried at Makrana, Rajasthan (~300-350 km transport distance). The marble consists of >98% pure calcite, rendering it translucent to sunlight. The surrounding gates, mosques, and plinths utilize red sandstone from Dholpur and Tantpur. Decorative elements incorporate intricate Pietra Dura stone inlay using semi-precious agate, carnelian, lapis lazuli, and jade."),
    
    ("Historical Centre of Arequipa (Peru)", "Volcanic Ignimbrite Tuff (Sillar de Arequipa)",
     "Arequipa's unique urban architecture is built almost entirely of 'Sillar'—a light, porous, white-to-pinkish dacitic-rhyolitic vitric tuff deposited by pyroclastic flows from nearby volcanoes (El Misti and Chachani). Quarried locally from the Añashuayco Ravine, Sillar blocks are easily carved into intricate Mestizo Baroque relief ornament while providing high thermal insulation and structural elasticity during seismic events."),
    
    ("Great Zimbabwe National Monument (Zimbabwe)", "Proterozoic Biotite Granite",
     "Constructed between the 11th and 15th centuries, Great Zimbabwe features massive dry-stone granite walls standing up to 11 meters high without any mortar. The builders utilized natural thermal exfoliation, harvesting exfoliated granite slabs from surrounding batholith outcrops. Sculptural uprights utilized soft talc-schist (steatite / soapstone), forming the famous carved Zimbabwe Birds."),
    
    ("Ellora Caves (Maharashtra, India)", "Deccan Trap Igneous Basalt",
     "Ellora represents monolithic rock-cut engineering carved directly into a steep basalt cliff of the Deccan Traps. Rather than assembling blocks, ancient masons excavated over 200,000 tonnes of solid basalt top-down to reveal the monolithic Kailash Temple (Cave 16). The fine-grained, dark grey basalt provided extraordinary compressive strength and structural stability."),
    
    ("Rapa Nui / Easter Island (Chile)", "Pyroclastic Volcanic Lava Tuff (Rano Raraku)",
     "The iconic Moai monolithic statues (887 recorded) were carved almost exclusively from yellow-brown volcanic lapilli tuff quarried inside the Rano Raraku crater. Red scoria from Puna Pau was used for topknot (Pukao) additions, while dark basalt was reserved for fine carving chisels. Anastylosis programs have systematically restored fallen moai onto coastal stone platforms (Ahu)."),
    
    ("Stonehenge & Avebury (Wiltshire, UK)", "Silicified Sandstone (Sarsen) & Welsh Bluestone",
     "Stonehenge incorporates two distinct stone types transported over substantial distances. The outer upright trilithons consist of local Sarsen (silicified Palaeogene sandstone) quarried ~25 km away on the Marlborough Downs. The inner circle consists of 'Bluestones' (dolerite, rhyolite, and volcanic tuffs) transported ~240 km from the Preseli Hills in Pembrokeshire, Wales."),
    
    ("Dholavira: A Harappan City (Gujarat, India)", "Calcareous Sandstone & Gemstone Assemblage",
     "Located on Khadir Island in the Rann of Kachchh, this 4,500-year-old Indus Valley metropolis was constructed using cut calcareous sandstone blocks extracted from nearby island quarries. Unlike mud-brick Harappan sites, Dholavira features dressed stone fortification walls, underground stone-lined water reservoirs, and specialized lapidary workshops processing agate, carnelian, and steatite beads."),
    
    ("Aachen Cathedral (Aachen, Germany)", "Devonian Fossiliferous Limestone & Imperial Spolia",
     "Charlemagne's Palatine Chapel (built 796–805 AD) combines local blue-grey Devonian limestone ('Aachener Blaustein') with ancient imperial spolia. Charlemagne legally secured antique marble and porphyry column shafts from ancient Roman structures in Rome and Ravenna, physically transporting them across the Alps to imbue his Carolingian capital with ancient imperial legitimacy.")
]

for title, mat, desc in case_studies:
    style_heading(doc.add_paragraph(), title, level=2)
    p_m = doc.add_paragraph()
    p_m.paragraph_format.space_before = Pt(0)
    p_m.paragraph_format.space_after = Pt(2)
    r_mk = p_m.add_run("Primary Material: ")
    r_mk.bold = True
    r_mk.font.name = 'Arial'
    r_mv = p_m.add_run(mat)
    r_mv.font.name = 'Arial'
    r_mv.font.color.rgb = RGBColor(43, 108, 176)
    
    add_body_p(doc, desc)

# Save document
output_filename = "UNESCO_World_Heritage_Stones_Analysis_Report.docx"
doc.save(output_filename)
print(f"Report 1 successfully saved to {output_filename}")
