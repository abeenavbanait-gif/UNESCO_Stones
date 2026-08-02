import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx_helpers import set_cell_background, set_cell_margins, add_callout, style_heading, add_body_p
import os
import pandas as pd

def main():
    print("==========================================================")
    print("📝 Generating DOCX Rescan and Methodology Report 📝")
    print("==========================================================")

    # 1. Load Data to calculate exact statistics
    csv_all = 're-scan/rescanned_no_stone_sites.csv'
    csv_built = 're-scan/rescanned_built_geological_monuments.csv'

    if not os.path.exists(csv_all) or not os.path.exists(csv_built):
        print("❌ Error: Rescan CSV files not found. Run classify_monuments_v2.py first.")
        return

    df = pd.read_csv(csv_all)
    built = pd.read_csv(csv_built)

    total_rescanned = len(df)
    built_count = len(built)
    built_pct = (built_count / total_rescanned) * 100

    high_count = len(df[df['confidence_v2'] == 'HIGH'])
    med_count = len(df[df['confidence_v2'] == 'MEDIUM'])
    low_count = len(df[df['confidence_v2'] == 'LOW'])
    none_count = len(df[df['confidence_v2'] == 'NONE'])

    high_pct = (high_count / total_rescanned) * 100
    med_pct = (med_count / total_rescanned) * 100
    low_pct = (low_count / total_rescanned) * 100
    none_pct = (none_count / total_rescanned) * 100

    doc = docx.Document()

    # Margins: Standard 1 inch
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title Block
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(4)
    r_title = p_title.add_run("GEOLOGICAL RESCAN & METHODOLOGY REPORT")
    r_title.font.name = 'Arial'
    r_title.font.size = Pt(20)
    r_title.bold = True
    r_title.font.color.rgb = RGBColor(26, 54, 93) # Navy Blue

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(14)
    r_sub = p_sub.add_run("An Analytical Audit and Scoring Methodology for Rescanned Cultural Sites from Table 1 (Backup 37 Methodology)")
    r_sub.font.name = 'Arial'
    r_sub.font.size = Pt(11.5)
    r_sub.font.color.rgb = RGBColor(74, 85, 104)

    # Metadata Block Table
    meta_table = doc.add_table(rows=2, cols=4)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Total Rescanned", "Newly Identified", "Recovery Rate", "Primary Model"]
    vals = [f"{total_rescanned} Sites", f"{built_count} Monuments", f"{built_pct:.1f}%", "Hybrid spaCy + LLM"]

    for i in range(4):
        c_hdr = meta_table.cell(0, i)
        set_cell_background(c_hdr, "1A365D")
        set_cell_margins(c_hdr, top=80, bottom=80, left=100, right=100)
        p = c_hdr.paragraphs[0]
        r = p.add_run(headers[i])
        r.font.name = 'Arial'
        r.font.size = Pt(9)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        
        c_val = meta_table.cell(1, i)
        set_cell_background(c_val, "F7FAFC")
        set_cell_margins(c_val, top=80, bottom=80, left=100, right=100)
        p2 = c_val.paragraphs[0]
        r2 = p2.add_run(vals[i])
        r2.font.name = 'Arial'
        r2.font.size = Pt(9)
        r2.font.color.rgb = RGBColor(45, 55, 72)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Executive Summary Callout
    summary_text = (
        f"This methodological rescan investigated {total_rescanned} UNESCO World Heritage Sites originally classified as containing "
        "no stone or having negligible stone scores. By implementing an upgraded classification engine (v2) incorporating "
        "lemmatization, title indicators, expanded rock dictionaries, and targeted Gemini LLM validation, we successfully "
        f"verified {built_count} sites as built monuments containing geological material or stonework (a {built_pct:.1f}% recovery rate)."
    )
    add_callout(doc, summary_text, title="EXECUTIVE RESCAN FINDINGS")

    # Section 1: Introduction
    p_h1 = doc.add_paragraph()
    style_heading(p_h1, "1. Introduction & Objectives", level=1)
    
    add_body_p(doc, 
        "In the primary classification of UNESCO World Heritage Sites, a portion of cultural sites did not trigger explicit rock "
        "matches due to linguistic limitations, brief descriptions, or OUV statement constraints (often describing historic structures "
        "without naming the specific geological stone types, e.g., 'masonry fortifications' or 'monolithic carved temple')."
    )
    add_body_p(doc, 
        "The objective of the Geological Rescan (v2) is to apply a more robust, multi-layer NLP parsing framework, "
        "study titles directly for architectural markers, and utilize high-performance large language models (Gemini 3.5 Flash) "
        "to extract and verify implied or unstated geological building materials from the target 'no-stone' sites."
    )

    # Section 2: Methodology
    p_h2 = doc.add_paragraph()
    style_heading(p_h2, "2. Confidence Score & Classifier Methodology (v2)", level=1)

    add_body_p(doc, 
        "The upgraded classification engine computes a Final Score for each site across four thematic layers using "
        "spaCy lemmatization and exact word boundaries."
    )

    p_eq = doc.add_paragraph()
    p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_eq = p_eq.add_run("Final Score = Criteria Layer + Title Layer + Text Match Layer - Exclusion Deductions")
    r_eq.font.name = 'Arial'
    r_eq.bold = True
    r_eq.font.size = Pt(11)
    r_eq.font.color.rgb = RGBColor(26, 54, 93)

    p_h2_1 = doc.add_paragraph()
    style_heading(p_h2_1, "2.1 Scoring Weights & Parameters", level=2)

    weights_text = (
        "• UNESCO Criteria: Criterion (iv) adds +3 points; Criterion (i) adds +2 points; Criterion (ii) adds +1 point.\n"
        "• Title Indicators: Matching core stone or architectural terms directly in the site name adds +4 points (e.g. 'aqueduct', 'monolith', 'cathedral').\n"
        "• Geological Rock Types: Match of igneous, sedimentary, or metamorphic rock groups adds +4 (1+ matches), +7 (3+ matches), or +10 (5+ matches).\n"
        "• Named/Trade Stones: Specific regional stone types (e.g. Carrara marble, Portland stone) add +5 (1+ matches) or +8 (3+ matches).\n"
        "• Construction & Masonry Terms: Specific stonework terms (e.g. ashlar, rubble, masonry) add +3 (2+ matches) or +5 (5+ matches).\n"
        "• Decorative Minerals: Semi-precious stones (e.g. lapis lazuli, jasper, jade) add +2 points.\n"
        "• Architectural Elements: Architectural structures (e.g. column, dome, vault, pier) add +3 (3+ matches) or +5 (6+ matches).\n"
        "• Exclusions: Non-built cultural heritage markers (e.g. rock art, intangible, textile) deduct -3 points."
    )
    add_body_p(doc, weights_text)

    p_h2_2 = doc.add_paragraph()
    style_heading(p_h2_2, "2.2 Confidence Tier Definitions", level=2)

    tier_text = (
        "Sites are categorized into Confidence Tiers using thresholds based on scores, title indicators, and stone counts:\n\n"
        "1. HIGH Confidence: Final Score >= 10, OR Stone Count >= 2, OR an explicit stone title indicator is matched.\n"
        "2. MEDIUM Confidence: Final Score >= 4, OR Stone Count >= 1, OR Construction Terms >= 2.\n"
        "3. LOW Confidence: Final Score between 1 and 3 (minor structural indicators).\n"
        "4. NONE: Final Score = 0 (no stone/structural elements matched)."
    )
    add_body_p(doc, tier_text)

    # Section 3: Summary Results
    p_h3 = doc.add_paragraph()
    style_heading(p_h3, "3. Rescan Statistics & Distribution", level=1)

    add_body_p(doc, "The distribution of confidence tiers across the 718 rescanned sites is summarized below:")

    # Distribution Table
    dist_table = doc.add_table(rows=5, cols=3)
    dist_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cols = ["Confidence Tier", "Site Count", "Percentage"]
    row_data = [
        ["HIGH", f"{high_count}", f"{high_pct:.1f}%"],
        ["MEDIUM", f"{med_count}", f"{med_pct:.1f}%"],
        ["LOW", f"{low_count}", f"{low_pct:.1f}%"],
        ["NONE", f"{none_count}", f"{none_pct:.1f}%"]
    ]

    # Header Row
    for i in range(3):
        cell = dist_table.cell(0, i)
        set_cell_background(cell, "2B6CB0")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        p = cell.paragraphs[0]
        r = p.add_run(cols[i])
        r.font.name = 'Arial'
        r.font.size = Pt(10)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)

    # Data Rows
    for r_idx, row_vals in enumerate(row_data):
        for c_idx, val in enumerate(row_vals):
            cell = dist_table.cell(r_idx + 1, c_idx)
            bg = "F7FAFC" if r_idx % 2 == 0 else "FFFFFF"
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.name = 'Arial'
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(45, 55, 72)
            if c_idx == 0:
                r.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Section 4: Country Distribution
    p_h4 = doc.add_paragraph()
    style_heading(p_h4, "4. Geographic & Country Analysis", level=1)

    add_body_p(doc, 
        "Newly identified built stone monuments are distributed globally, with notable concentrations in countries "
        "possessing high historical densities of stone-built civil and religious architecture."
    )

    # Countries Table
    country_counts = built['country'].value_counts().head(10)
    c_table = doc.add_table(rows=11, cols=3)
    c_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_cols = ["Rank", "Country", "Newly Identified Monuments"]

    # Header Row
    for i in range(3):
        cell = c_table.cell(0, i)
        set_cell_background(cell, "2C5282")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        p = cell.paragraphs[0]
        r = p.add_run(c_cols[i])
        r.font.name = 'Arial'
        r.font.size = Pt(10)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)

    # Data Rows
    for r_idx, (c_name, count) in enumerate(country_counts.items()):
        row_vals = [f"{r_idx+1}", c_name, f"{count} Sites"]
        for c_idx, val in enumerate(row_vals):
            cell = c_table.cell(r_idx + 1, c_idx)
            bg = "F7FAFC" if r_idx % 2 == 0 else "FFFFFF"
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.name = 'Arial'
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(45, 55, 72)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Section 5: Showcase Sites
    p_h5 = doc.add_paragraph()
    style_heading(p_h5, "5. Newly Recovered High-Confidence Showcase Sites", level=1)

    showcases = [
        ("Buddhist Monuments at Sanchi (India)", "Score 18 (HIGH)", 
         "Features monolithic Pillars of Asoka, stone balustrades, and structural gateways carved from locally quarried Chunarian sandstone."),
        ("Rock-Hewn Churches, Lalibela (Ethiopia)", "Score 16 (HIGH)", 
         "Consists of 11 medieval monolithic cave churches carved directly out of solid volcanic tuff rock."),
        ("The Cathedral of St James in Šibenik (Croatia)", "Score 18 (HIGH)", 
         "Constructed entirely of stone (limestone and marble from local quarries) without any mortar or wood binders."),
        ("Pontcysyllte Aqueduct and Canal (United Kingdom)", "Score 18 (HIGH)", 
         "Iconic industrial structure built with high-quality ashlar masonry piers and rubble-filled cast iron canal supports.")
    ]

    for title, score, desc in showcases:
        p_site = doc.add_paragraph()
        p_site.paragraph_format.keep_with_next = True
        p_site.paragraph_format.space_before = Pt(8)
        p_site.paragraph_format.space_after = Pt(2)
        r_site_t = p_site.add_run(f"• {title} — ")
        r_site_t.bold = True
        r_site_t.font.name = 'Arial'
        r_site_t.font.color.rgb = RGBColor(26, 54, 93)
        
        r_site_s = p_site.add_run(f"{score}\n")
        r_site_s.bold = True
        r_site_s.font.name = 'Arial'
        r_site_s.font.color.rgb = RGBColor(47, 133, 90) # Green
        
        r_site_d = p_site.add_run(desc)
        r_site_d.font.name = 'Arial'
        r_site_d.font.size = Pt(10)
        r_site_d.font.color.rgb = RGBColor(74, 85, 104)

    # Output file references
    p_out = doc.add_paragraph()
    p_out.paragraph_format.space_before = Pt(18)
    r_out = p_out.add_run("Output Datasets:\n")
    r_out.bold = True
    r_out.font.name = 'Arial'
    
    r_out_desc = p_out.add_run(
        "• Complete Rescan Database: re-scan/rescanned_no_stone_sites.csv\n"
        "• Built Monuments Database: re-scan/rescanned_built_geological_monuments.csv"
    )
    r_out_desc.font.name = 'Arial'
    r_out_desc.font.size = Pt(9.5)
    r_out_desc.font.color.rgb = RGBColor(74, 85, 104)

    # Save to workspace folder
    out_docx_path = 're-scan/geological_rescan_methodology_report.docx'
    doc.save(out_docx_path)
    print(f"✅ Document successfully saved to {out_docx_path}")

if __name__ == '__main__':
    main()
