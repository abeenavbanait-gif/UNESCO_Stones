import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx_helpers import set_cell_background, set_cell_margins, add_callout, style_heading, add_body_p
import os

doc = docx.Document()

# Page setup: Standard 1 inch margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ==============================================================================
# TITLE BLOCK
# ==============================================================================
p_title = doc.add_paragraph()
p_title.paragraph_format.space_before = Pt(0)
p_title.paragraph_format.space_after = Pt(4)
r_title = p_title.add_run("GEOLOGICAL BIOGRAPHY & METHODOLOGICAL MONOGRAPH OF WORLD HERITAGE")
r_title.font.name = 'Arial'
r_title.font.size = Pt(22)
r_title.bold = True
r_title.font.color.rgb = RGBColor(26, 54, 93)

p_sub = doc.add_paragraph()
p_sub.paragraph_format.space_after = Pt(14)
r_sub = p_sub.add_run("A Master Empirical Investigation into Lithology, Sourcing Provenance, Dataset Selection Protocol, Operational Barriers, and Conservation Dynamics Across UNESCO World Heritage Sites (Backup 37 Master Dataset)")
r_sub.font.name = 'Arial'
r_sub.font.size = Pt(11.5)
r_sub.font.color.rgb = RGBColor(74, 85, 104)

# Metadata Table
meta_table = doc.add_table(rows=2, cols=4)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["Master Dataset", "Record Universe", "Skipped Research Queue", "OUV Absent / Inaccessible"]
vals = ["Backup 37 (July 2026)", "902 Sites / 80 Attributes", "650 Sites (72.1%)", "49 Sites (5.4%)"]

for i in range(4):
    c_hdr = meta_table.cell(0, i)
    set_cell_background(c_hdr, "1A365D")
    set_cell_margins(c_hdr, top=80, bottom=80, left=100, right=100)
    p = c_hdr.paragraphs[0]
    r = p.add_run(headers[i])
    r.font.name = 'Arial'
    r.font.size = Pt(9)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)
    
    c_val = meta_table.cell(1, i)
    set_cell_background(c_val, "F7FAFC")
    set_cell_margins(c_val, top=80, bottom=80, left=100, right=100)
    p2 = c_val.paragraphs[0]
    r2 = p2.add_run(vals[i])
    r2.font.name = 'Arial'
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(45, 55, 72)

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# ==============================================================================
# SECTION 1: EXECUTIVE SUMMARY & CORE EMPIRICAL BASELINE
# ==============================================================================
style_heading(doc.add_paragraph(), "1. Executive Summary & Core Empirical Baseline", level=1)

add_body_p(doc, "The UNESCO World Heritage list represents the pinnacle of human architectural, artistic, and historical accomplishment. Yet every standing monument, ancient citadel, rock-cut sanctuary, and cathedral is fundamentally a geological artifact—constructed from specific lithologies extracted from local or regional Earth formations. This master monograph synthesizes the empirical findings from the complete Heritage Stones dataset (Backup 37), encompassing 902 World Heritage Sites across 168 sovereign nations evaluated across 80 distinct structural, lithological, provenance, and conservation attributes.")

add_callout(doc, 
    "Out of 902 target built-monument sites, exactly 193 entries (21.4%) contain verified primary stone identification data following the Backup 37 quality audit. Among the 129 provenance-classified monuments, 88.4% (114 sites) relied exclusively on local stone extraction within a 5 to 25 km radius. Sedimentary formations (limestone and sandstone) account for 60.5% of all classified stone heritage, reflecting both the historical workability of carbonate rocks under manual masonry tools and their widespread geographical occurrence near historic urban centers.",
    title="EXECUTIVE MONOGRAPH SUMMARY", border_color="1A365D", bg_color="F7FAFC")

# Key Metrics Table
t_metrics = doc.add_table(rows=7, cols=3)
t_metrics.alignment = WD_TABLE_ALIGNMENT.CENTER
t_headers = ["Analytical Dimension", "Empirical Baseline (Backup 37)", "Primary Heritage Significance"]
t_data = [
    ["Researched Stone Sites", "193 / 902 (21.4%)", "Verified research baseline; 709 sites form the unstudied expansion queue."],
    ["Dominant Rock Class", "Sedimentary (60.5%)", "Limestone (24 sites) and Sandstone (14 sites) anchor global masonry."],
    ["Resource Provenance", "Local (88.4%)", "114 of 129 provenance-classified sites extracted stone within immediate vicinity."],
    ["Strategic Sourcing (Imports)", "Imported (4.7%)", "Long-distance trade reserved for imperial prestige (Taj Mahal, CST Mumbai)."],
    ["Preservation Baseline", "78.4% Good/Excellent", "152 of 194 assessed sites exhibit strong structural integrity under UNESCO oversight."],
    ["Primary Degradation Vector", "Surface Erosion (31 sites)", "Rainwater leaching, wind abrasion, and humidity dominate structural threats."]
]

for col_idx, h in enumerate(t_headers):
    c = t_metrics.cell(0, col_idx)
    set_cell_background(c, "1A365D")
    set_cell_margins(c, top=80, bottom=80, left=100, right=100)
    p = c.paragraphs[0]
    r = p.add_run(h)
    r.font.name = 'Arial'
    r.font.size = Pt(9.5)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)

for row_idx, r_data in enumerate(t_data):
    bg = "FFFFFF" if row_idx % 2 == 0 else "F7FAFC"
    for col_idx, val in enumerate(r_data):
        c = t_metrics.cell(row_idx + 1, col_idx)
        set_cell_background(c, bg)
        set_cell_margins(c, top=70, bottom=70, left=100, right=100)
        p = c.paragraphs[0]
        r = p.add_run(val)
        r.font.name = 'Arial'
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(45, 55, 72)
        if col_idx == 1:
            r.bold = True

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# ==============================================================================
# SECTION 2: METHODOLOGICAL RATIONALE — DERIVING THE 902 SITES FROM 972 CANDIDATES
# ==============================================================================
style_heading(doc.add_paragraph(), "2. Methodological Rationale: Deriving the 902 Dataset from 972 Candidates", level=1)

add_body_p(doc, "A foundational requirement of rigorous heritage data science is establishing an un-biased, reproducible target universe. When constructing a global inventory of heritage stones, assessing built monuments requires isolating sites that contain physical architecture, masonry, rock-cut excavations, or structural earth/stone materials from sites inscribed purely for natural ecosystems, agricultural practices, or intangible social memories.")

add_body_p(doc, "Evaluating the UNESCO World Heritage candidate pool revealed 972 candidate Cultural Sites inscribed under Criteria (i) through (vi). A 4-stage screening protocol filtered out 70 sites whose primary Outstanding Universal Value (OUV) is non-structural, yielding the 902 target built-monument universe (92.8% inclusion rate):")

# Image: Exclusion Donut & Stages Bar
if os.path.exists('charts/exclusion_donut.png'):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after = Pt(4)
    p_img.add_run().add_picture('charts/exclusion_donut.png', width=Inches(5.0))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(10)
    r_cap = p_cap.add_run("Figure 1: Ratio of Target Built-Monument Universe (n=902) to Excluded Non-Structural Sites (n=70).")
    r_cap.font.name = 'Arial'
    r_cap.font.size = Pt(8.5)
    r_cap.italic = True
    r_cap.font.color.rgb = RGBColor(113, 128, 150)

# Summary table of 4 stages
t_stages = doc.add_table(rows=5, cols=4)
t_stages.alignment = WD_TABLE_ALIGNMENT.CENTER
t_s_headers = ["Screening Stage", "Exclusion Category", "Sites Excluded", "Primary Methodological Rationale"]
t_s_data = [
    ["Stage 1", "Non-Structural Cultural Landscapes", "32 Sites", "Agro-ecological farming, vineyard terraces, and botanical domains lacking built masonry monuments."],
    ["Stage 2", "Pure Timber & Organic Architecture", "21 Sites", "Timber framing, log joinery, reed thatch without stone foundations or ashlar cladding."],
    ["Stage 3", "Submerged & Maritime Archaeology", "10 Sites", "Submerged lakebed stilt settlements and underwater shipwreck reserves."],
    ["Stage 4", "Intangible Social Memory Zones", "7 Sites", "Commemorative memorial landscapes inscribed for symbolic event memory under Criterion (vi)."]
]

for col_idx, h in enumerate(t_s_headers):
    c = t_stages.cell(0, col_idx)
    set_cell_background(c, "1A365D")
    set_cell_margins(c, top=80, bottom=80, left=100, right=100)
    p = c.paragraphs[0]
    r = p.add_run(h)
    r.font.name = 'Arial'
    r.font.size = Pt(9.5)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)

for row_idx, r_data in enumerate(t_s_data):
    bg = "FFFFFF" if row_idx % 2 == 0 else "F7FAFC"
    for col_idx, val in enumerate(r_data):
        c = t_stages.cell(row_idx + 1, col_idx)
        set_cell_background(c, bg)
        set_cell_margins(c, top=70, bottom=70, left=100, right=100)
        p = c.paragraphs[0]
        r = p.add_run(val)
        r.font.name = 'Arial'
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(45, 55, 72)
        if col_idx == 2:
            r.bold = True

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# ==============================================================================
# SECTION 3: DATASET ARCHITECTURE & FIELD COMPLETENESS
# ==============================================================================
style_heading(doc.add_paragraph(), "3. Dataset Architecture & Field Completeness Baseline", level=1)

add_body_p(doc, "The 80 attributes in the dataset track three core operational dimensions: administrative metadata (Criteria, Country, Architecture Type), geological petrographic data (Rock Class, Lithology, Formation, Minerals, Colour, Texture), and architectural conservation metrics (Masonry Technique, Structural Use, Quarry Origin, Weathering, Condition, Restoration).")

# Field completeness table
t_fc = doc.add_table(rows=16, cols=3)
t_fc.alignment = WD_TABLE_ALIGNMENT.CENTER
t_fc_headers = ["Attribute Field Name", "Filled Records (n=902)", "Completion Percentage (%)"]
t_fc_data = [
    ["Site ID / Site Name / Country", "902", "100.0%"],
    ["Architecture Type", "901", "99.9%"],
    ["UNESCO Criteria", "795", "88.1%"],
    ["Site Condition Rating", "194", "21.5%"],
    ["Mentioned Major Stone(s)", "193", "21.4%"],
    ["Structural Use Application", "190", "21.1%"],
    ["Restoration Approach", "174", "19.3%"],
    ["Masonry Technique", "157", "17.4%"],
    ["Rock Class Classification", "152", "16.9%"],
    ["Local vs Imported Provenance", "129", "14.3%"],
    ["Weathering Threat Vectors", "119", "13.2%"],
    ["Lithological Description", "80", "8.9%"],
    ["Quarry Location / Country", "62", "6.9%"],
    ["Local Vernacular Stone Name", "43", "4.8%"],
    ["Mineralogical Composition", "22", "2.4%"]
]

for col_idx, h in enumerate(t_fc_headers):
    c = t_fc.cell(0, col_idx)
    set_cell_background(c, "1A365D")
    set_cell_margins(c, top=70, bottom=70, left=100, right=100)
    p = c.paragraphs[0]
    r = p.add_run(h)
    r.font.name = 'Arial'
    r.font.size = Pt(9.5)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)

for row_idx, r_data in enumerate(t_fc_data):
    bg = "FFFFFF" if row_idx % 2 == 0 else "F7FAFC"
    for col_idx, val in enumerate(r_data):
        c = t_fc.cell(row_idx + 1, col_idx)
        set_cell_background(c, bg)
        set_cell_margins(c, top=60, bottom=60, left=100, right=100)
        p = c.paragraphs[0]
        r = p.add_run(val)
        r.font.name = 'Arial'
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(45, 55, 72)
        if col_idx == 2:
            r.bold = True

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# Add Image 5: Data Completeness Bar
if os.path.exists('charts/data_completeness_bar.png'):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after = Pt(4)
    p_img.add_run().add_picture('charts/data_completeness_bar.png', width=Inches(6.0))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(14)
    r_cap = p_cap.add_run("Figure 2: Attribute completeness percentages across 902 World Heritage Sites (Backup 37).")
    r_cap.font.name = 'Arial'
    r_cap.font.size = Pt(8.5)
    r_cap.italic = True
    r_cap.font.color.rgb = RGBColor(113, 128, 150)

# ==============================================================================
# SECTION 4: GEOLOGICAL CLASSIFICATION & CONTINENTAL DISTRIBUTIONS
# ==============================================================================
style_heading(doc.add_paragraph(), "4. Geological Classification: The Three Kingdoms of Stone", level=1)

add_body_p(doc, "Of the 152 sites with confirmed rock class entries, sedimentary geology dominates world heritage construction. Sedimentary rocks comprise 60.5% (92 sites), followed by igneous formations at 24.3% (37 sites), and metamorphic rocks at 15.1% (23 sites).")

# Add Image 1: Rock Class Donut
if os.path.exists('charts/rock_class_donut.png'):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after = Pt(4)
    p_img.add_run().add_picture('charts/rock_class_donut.png', width=Inches(5.2))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(14)
    r_cap = p_cap.add_run("Figure 3: Distribution of primary rock classes across classified heritage monuments (n=152).")
    r_cap.font.name = 'Arial'
    r_cap.font.size = Pt(8.5)
    r_cap.italic = True
    r_cap.font.color.rgb = RGBColor(113, 128, 150)

add_body_p(doc, "Sedimentary dominance is driven by three geological and historical realities: (1) superior manual workability of carbonate rocks under hand tools, (2) widespread occurrence of shallow marine carbonate platforms across Southern Europe, North Africa, and South Asia where urban civilization developed, and (3) natural river valley bluff exposures minimizing overburden mining costs.")

style_heading(doc.add_paragraph(), "Regional Continental Patterns", level=2)
add_body_p(doc, "Europe leads total classified sites (66 sites) with heavy reliance on sedimentary limestone cathedrals and Mediterranean marble architecture. Asia (39 sites) exhibits a dual reliance on sedimentary sandstones (Harappan and Mughal sites) and Deccan igneous basalts (rock-cut cave temples). Africa and the Americas display balanced igneous-sedimentary distributions, utilizing volcanic tuffs and granites.")

# Add Image 4: Regional Rock Class
if os.path.exists('charts/regional_rock_class.png'):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after = Pt(4)
    p_img.add_run().add_picture('charts/regional_rock_class.png', width=Inches(5.8))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(14)
    r_cap = p_cap.add_run("Figure 4: Rock class distribution broken down across geographic regions.")
    r_cap.font.name = 'Arial'
    r_cap.font.size = Pt(8.5)
    r_cap.italic = True
    r_cap.font.color.rgb = RGBColor(113, 128, 150)

# ==============================================================================
# SECTION 5: TYPOLOGICAL DOMINANCE & MATERIAL TAXONOMY
# ==============================================================================
style_heading(doc.add_paragraph(), "5. Typological Dominance & Material Taxonomy", level=1)

add_body_p(doc, "Evaluating specific stone types identifies Limestone as the single most common major stone in world heritage (24 direct mentions), followed by Granite (15 entries), Sandstone (14 entries), and Marble (10 entries). Following the Backup 37 quality audit, spurious 'Flint' assignments across 6 non-prehistoric sites were corrected, leaving Spiennes (Belgium) as the genuine Neolithic flint mining entry.")

# Add Image 2: Top Stones Bar Chart
if os.path.exists('charts/top_stones_bar.png'):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after = Pt(4)
    p_img.add_run().add_picture('charts/top_stones_bar.png', width=Inches(5.8))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(14)
    r_cap = p_cap.add_run("Figure 5: Frequency of primary stone types directly identified across UNESCO inscriptions (Backup 37 Audit).")
    r_cap.font.name = 'Arial'
    r_cap.font.size = Pt(8.5)
    r_cap.italic = True
    r_cap.font.color.rgb = RGBColor(113, 128, 150)

add_body_p(doc, "The dataset also captures highly specific local vernacular stone names: Sillar (Peru dacitic ignimbrite), Makrana Marble (India >98% calcitic marble), Pietra Serena and Pietra Forte (Florentine sandstones), Pietra d'Istria (Venetian marine-resistant limestone), Aachener Blaustein (Germany), Kabook (Sri Lanka), Daga (Zimbabwe), and Sarsen (UK).")

# ==============================================================================
# SECTION 6: PROVENANCE DYNAMICS & LOGISTICS
# ==============================================================================
style_heading(doc.add_paragraph(), "6. Provenance Dynamics: Local Sourcing vs. Strategic Imports", level=1)

add_body_p(doc, "Among 129 provenance-classified sites, 114 sites (88.4%) relied entirely on local stone extraction within a 5 to 25 km radius. Only 6 sites (4.7%) relied exclusively on long-distance imports, while 8 sites (6.2%) combined local structural stone with imported decorative elements.")

# Add Image 3: Provenance Pie
if os.path.exists('charts/provenance_pie.png'):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after = Pt(4)
    p_img.add_run().add_picture('charts/provenance_pie.png', width=Inches(5.2))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(14)
    r_cap = p_cap.add_run("Figure 6: Sourcing provenance distribution across documented heritage sites (n=129).")
    r_cap.font.name = 'Arial'
    r_cap.font.size = Pt(8.5)
    r_cap.italic = True
    r_cap.font.color.rgb = RGBColor(113, 128, 150)

add_body_p(doc, "Pre-industrial overland transport of heavy stone blocks imposed immense costs. Long-distance stone transport occurred almost exclusively via maritime/river shipping routes (Istrian stone shipped to Venice) or royal imperial decree where cost was secondary to state symbolism (Makrana marble hauled 300+ km to Agra for the Taj Mahal, or Welsh bluestones moved 240 km to Stonehenge).")

# ==============================================================================
# SECTION 7: PRESERVATION DYNAMICS & WEATHERING THREATS
# ==============================================================================
style_heading(doc.add_paragraph(), "7. Preservation Dynamics & Degradation Vectors", level=1)

add_body_p(doc, "Across 194 assessed sites, 78.4% (152 sites) are rated in Good or Excellent condition, testifying to active UNESCO management. However, 16.5% (32 sites) exhibit moderate degradation, and 5.2% (10 sites) suffer from poor or critical condition ratings requiring immediate conservation intervention.")

# Add Image 6: Condition Donut
if os.path.exists('charts/condition_assessment_pie.png'):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after = Pt(4)
    p_img.add_run().add_picture('charts/condition_assessment_pie.png', width=Inches(5.2))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(14)
    r_cap = p_cap.add_run("Figure 7: Environmental condition assessment distribution across evaluated monuments (n=194).")
    r_cap.font.name = 'Arial'
    r_cap.font.size = Pt(8.5)
    r_cap.italic = True
    r_cap.font.color.rgb = RGBColor(113, 128, 150)

add_body_p(doc, "Evaluating weathering threats across 119 sites identifies Surface Erosion (31 sites) and High Ambient Humidity (17 sites) as primary agents of material decay. Atmospheric pollution and sulfur dioxide emissions (11 sites) drive severe sulfation of calcium carbonate in limestones and marbles into soluble gypsum crusts.")

# Add Image 7: Weathering Threats Bar
if os.path.exists('charts/weathering_threats_bar.png'):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after = Pt(4)
    p_img.add_run().add_picture('charts/weathering_threats_bar.png', width=Inches(5.8))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(14)
    r_cap = p_cap.add_run("Figure 8: Frequency of primary weathering vectors affecting heritage stone structures.")
    r_cap.font.name = 'Arial'
    r_cap.font.size = Pt(8.5)
    r_cap.italic = True
    r_cap.font.color.rgb = RGBColor(113, 128, 150)

# ==============================================================================
# SECTION 8: OPERATIONAL BARRIERS TO STUDY: SKIPPED QUEUE & OUV ABSENT RECORDS
# ==============================================================================
style_heading(doc.add_paragraph(), "8. Operational Barriers to Study: 650 Skipped Sites & 49 OUV Absent Records", level=1)

add_body_p(doc, "A critical dimension of this research is understanding why 709 out of 902 built-monument sites (78.6%) currently remain unstudied in the primary dataset. The empirical evidence demonstrates that two operational barriers prevent immediate lithological extraction:")

add_callout(doc,
    "OPERATIONAL DATA BARRIERS:\n"
    "1. SKIPPED RESEARCH QUEUE (650 Sites | 72.1%): Sites where initial summary review confirmed built heritage, but stone extraction was deferred to Phase 2 due to finite researcher bandwidth (~3-15 hours required per site to consult multi-page ICOMOS evaluations and academic journals).\n"
    "2. OUV ABSENT / INACCESSIBLE (49 Sites | 5.4%): Sites where the official UNESCO Statement of Outstanding Universal Value (OUV) is completely missing, unindexed, or omitted from UNESCO's digital API, creating a digital 404 wall.\n"
    "3. OUV TEXT ISSUES (22 Sites | 2.4%): Sites where official OUV text exists but completely ignores building materials in favor of religious or historical narratives.",
    title="DOCUMENTATION BARRIERS AUDIT", border_color="C53030", bg_color="FFF5F5")

# Table of Operational Barriers
t_barriers = doc.add_table(rows=5, cols=4)
t_barriers.alignment = WD_TABLE_ALIGNMENT.CENTER
t_b_headers = ["Operational Category", "Site Count", "% of Dataset (n=902)", "Primary Mechanism Preventing Study"]
t_b_data = [
    ["Skipped Research Queue", "650 Sites", "72.1%", "Archival extraction deferred; requires manual search across external peer journals and ICOMOS reports."],
    ["OUV Text Inaccessible (Absent)", "49 Sites", "5.4%", "Official UNESCO OUV documentation is unindexed or missing from digital APIs; forces non-digital hardcopy archival research."],
    ["OUV Text Material Silence (Issue)", "22 Sites", "2.4%", "OUV text exists but focuses purely on historical/cultural significance, omitting physical stone names."],
    ["Researched & Verified Baseline", "193 Sites", "21.4%", "Active verified baseline with complete petrographic citations."]
]

for col_idx, h in enumerate(t_b_headers):
    c = t_barriers.cell(0, col_idx)
    set_cell_background(c, "742A2A")
    set_cell_margins(c, top=80, bottom=80, left=100, right=100)
    p = c.paragraphs[0]
    r = p.add_run(h)
    r.font.name = 'Arial'
    r.font.size = Pt(9.5)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)

for row_idx, r_data in enumerate(t_b_data):
    bg = "FFFFFF" if row_idx % 2 == 0 else "FFF5F5"
    for col_idx, val in enumerate(r_data):
        c = t_barriers.cell(row_idx + 1, col_idx)
        set_cell_background(c, bg)
        set_cell_margins(c, top=70, bottom=70, left=100, right=100)
        p = c.paragraphs[0]
        r = p.add_run(val)
        r.font.name = 'Arial'
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(45, 55, 72)
        if col_idx == 1:
            r.bold = True

doc.add_paragraph().paragraph_format.space_after = Pt(12)

style_heading(doc.add_paragraph(), "Detailed Audit of the 49 OUV Absent Sites", level=2)
add_body_p(doc, "The 49 'OUV Absent' sites represent a critical digital infrastructure gap. Because UNESCO's central digital database lacks indexed OUV text for these records, automated scrapers and standard archival search protocols hit a dead end. Researchers are forced to hunt down physical hardcopy dossiers, non-digitized PDF nomination files from national ministries, or regional heritage monographs. Landmark monuments trapped in this category include:")

# Table of sample OUV Absent sites
t_absent = doc.add_table(rows=7, cols=3)
t_absent.alignment = WD_TABLE_ALIGNMENT.CENTER
t_a_headers = ["OUV Absent Monument", "Country", "Expected Physical Lithology & Heritage Importance"]
t_a_data = [
    ["Old City of Dubrovnik", "Croatia", "Korčula and Brač Cretaceous limestone masonry fortifications and stone paving."],
    ["Selimiye Mosque Complex", "Türkiye", "Marmara marble columns, local limestone masonry, and granite piers."],
    ["Ephesus", "Türkiye", "Hellenistic and Roman marble temples, agoras, and paved avenues."],
    ["Historical Complex of Split", "Croatia", "Diocletian's Palace: Proconnesian marble and local Cretaceous limestone."],
    ["Rock Shelters of Bhimbetka", "India", "Vindhyan Sandstone quartzite cliff shelters featuring prehistoric rock art."],
    ["Murujuga Cultural Landscape", "Australia", "Precambrian granite batholiths featuring over 1 million petroglyphs."]
]

for col_idx, h in enumerate(t_a_headers):
    c = t_absent.cell(0, col_idx)
    set_cell_background(c, "742A2A")
    set_cell_margins(c, top=80, bottom=80, left=100, right=100)
    p = c.paragraphs[0]
    r = p.add_run(h)
    r.font.name = 'Arial'
    r.font.size = Pt(9.5)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)

for row_idx, r_data in enumerate(t_a_data):
    bg = "FFFFFF" if row_idx % 2 == 0 else "FFF5F5"
    for col_idx, val in enumerate(r_data):
        c = t_absent.cell(row_idx + 1, col_idx)
        set_cell_background(c, bg)
        set_cell_margins(c, top=70, bottom=70, left=100, right=100)
        p = c.paragraphs[0]
        r = p.add_run(val)
        r.font.name = 'Arial'
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(45, 55, 72)
        if col_idx == 0:
            r.bold = True

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# Add Image 8: Dataset Gap Donut
if os.path.exists('charts/data_gap_donut.png'):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after = Pt(4)
    p_img.add_run().add_picture('charts/data_gap_donut.png', width=Inches(5.2))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(14)
    r_cap = p_cap.add_run("Figure 9: Inscription status breakdown across the 902 World Heritage Sites universe (Backup 37 Audit).")
    r_cap.font.name = 'Arial'
    r_cap.font.size = Pt(8.5)
    r_cap.italic = True
    r_cap.font.color.rgb = RGBColor(113, 128, 150)

# Add Image 9: Country Unstudied Gaps Bar
if os.path.exists('charts/country_unstudied_gaps.png'):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after = Pt(4)
    p_img.add_run().add_picture('charts/country_unstudied_gaps.png', width=Inches(6.0))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(14)
    r_cap = p_cap.add_run("Figure 10: Unstudied site research gaps across leading World Heritage nations.")
    r_cap.font.name = 'Arial'
    r_cap.font.size = Pt(8.5)
    r_cap.italic = True
    r_cap.font.color.rgb = RGBColor(113, 128, 150)

# ==============================================================================
# SECTION 9: THE ICEBERG MODEL & BENCHMARK CASE STUDIES
# ==============================================================================
style_heading(doc.add_paragraph(), "9. The Iceberg Model & Benchmark Case Studies", level=1)

add_body_p(doc, "The current dataset represents an 'Iceberg Model': the 193 verified stone sites form the visible tip above water, while an estimated 532 unstudied stone monuments lie beneath the surface, hidden by documentation gaps.")

# Add Image 10: Iceberg Model Bar
if os.path.exists('charts/iceberg_uncaptured_model.png'):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after = Pt(4)
    p_img.add_run().add_picture('charts/iceberg_uncaptured_model.png', width=Inches(5.5))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(14)
    r_cap = p_cap.add_run("Figure 11: The Iceberg Model comparing verified baseline sites against estimated hidden stone mass.")
    r_cap.font.name = 'Arial'
    r_cap.font.size = Pt(8.5)
    r_cap.italic = True
    r_cap.font.color.rgb = RGBColor(113, 128, 150)

style_heading(doc.add_paragraph(), "Benchmark Case Studies", level=2)
case_studies = [
    ("Taj Mahal (Agra, India)", "Metamorphic Calcitic Marble (Makrana Marble) & Red Sandstone", 
     "Comprising 11 filled stone fields, the Taj Mahal represents the pinnacle of dual-material imperial masonry. The main mausoleum is clad in pure granoblastic calcitic marble quarried at Makrana, Rajasthan (~300-350 km transport distance). The marble consists of >98% pure calcite, rendering it translucent to sunlight. The surrounding gates, mosques, and plinths utilize red sandstone from Dholpur and Tantpur. Decorative elements incorporate intricate Pietra Dura stone inlay using semi-precious agate, carnelian, lapis lazuli, and jade."),
    
    ("Historical Centre of Arequipa (Peru)", "Volcanic Ignimbrite Tuff (Sillar de Arequipa)",
     "Arequipa's unique urban architecture is built almost entirely of 'Sillar'—a light, porous, white-to-pinkish dacitic-rhyolitic vitric tuff deposited by pyroclastic flows from nearby volcanoes (El Misti and Chachani). Quarried locally from the Añashuayco Ravine, Sillar blocks are easily carved into intricate Mestizo Baroque relief ornament while providing high thermal insulation and structural elasticity during seismic events."),
    
    ("Great Zimbabwe National Monument (Zimbabwe)", "Proterozoic Biotite Granite",
     "Constructed between the 11th and 15th centuries, Great Zimbabwe features massive dry-stone granite walls standing up to 11 meters high without any mortar. The builders utilized natural thermal exfoliation, harvesting exfoliated granite slabs from surrounding batholith outcrops. Sculptural uprights utilized soft talc-schist (steatite / soapstone), forming the famous carved Zimbabwe Birds."),
    
    ("Ellora Caves (Maharashtra, India)", "Deccan Trap Igneous Basalt",
     "Ellora represents monolithic rock-cut engineering carved directly into a steep basalt cliff of the Deccan Traps. Rather than assembling blocks, ancient masons excavated over 200,000 tonnes of solid basalt top-down to reveal the monolithic Kailash Temple (Cave 16). The fine-grained, dark grey basalt provided extraordinary compressive strength and structural stability."),
    
    ("Rapa Nui / Easter Island (Chile)", "Pyroclastic Volcanic Lava Tuff (Rano Raraku)",
     "The iconic Moai monolithic statues (887 recorded) were carved almost exclusively from yellow-brown volcanic lapilli tuff quarried inside the Rano Raraku crater. Red scoria from Puna Pau was used for topknot (Pukao) additions, while dark basalt was reserved for fine carving chisels. Anastylosis programs have systematically restored fallen moai onto coastal stone platforms (Ahu)."),
    
    ("Stonehenge & Avebury (Wiltshire, UK)", "Silicified Sandstone (Sarsen) & Welsh Bluestone",
     "Stonehenge incorporates two distinct stone types transported over substantial distances. The outer upright trilithons consist of local Sarsen (silicified Palaeogene sandstone) quarried ~25 km away on the Marlborough Downs. The inner circle consists of 'Bluestones' (dolerite, rhyolite, and volcanic tuffs) transported ~240 km from the Preseli Hills in Pembrokeshire, Wales."),
    
    ("Dholavira: A Harappan City (Gujarat, India)", "Calcareous Sandstone & Gemstone Assemblage",
     "Located on Khadir Island in the Rann of Kachchh, this 4,500-year-old Indus Valley metropolis was constructed using cut calcareous sandstone blocks extracted from nearby island quarries. Unlike mud-brick Harappan sites, Dholavira features dressed stone fortification walls, underground stone-lined water reservoirs, and specialized lapidary workshops processing agate, carnelian, and steatite beads."),
    
    ("Aachen Cathedral (Aachen, Germany)", "Devonian Fossiliferous Limestone & Imperial Spolia",
     "Charlemagne's Palatine Chapel (built 796–805 AD) combines local blue-grey Devonian limestone ('Aachener Blaustein') with ancient imperial spolia. Charlemagne legally secured antique marble and porphyry column shafts from ancient Roman structures in Rome and Ravenna, physically transporting them across the Alps to imbue his Carolingian capital with ancient imperial legitimacy.")
]

for title, mat, desc in case_studies:
    style_heading(doc.add_paragraph(), title, level=3)
    p_m = doc.add_paragraph()
    p_m.paragraph_format.space_before = Pt(0)
    p_m.paragraph_format.space_after = Pt(2)
    r_mk = p_m.add_run("Primary Material: ")
    r_mk.bold = True
    r_mk.font.name = 'Arial'
    r_mv = p_m.add_run(mat)
    r_mv.font.name = 'Arial'
    r_mv.font.color.rgb = RGBColor(43, 108, 176)
    
    add_body_p(doc, desc)

# Save master docx file
output_master_docx = "UNESCO_World_Heritage_Stones_Master_Comprehensive_Report_Backup37.docx"
doc.save(output_master_docx)
print(f"Master Monograph docx saved successfully to {output_master_docx}")

